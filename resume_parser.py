# # # import PyPDF2
# # # from docx import Document
# # # import re
# # # import os

# # # def extract_text_from_pdf(pdf_path):
# # #     """Extract text from PDF file"""
# # #     text = ""
# # #     try:
# # #         with open(pdf_path, 'rb') as file:
# # #             pdf_reader = PyPDF2.PdfReader(file)
# # #             for page in pdf_reader.pages:
# # #                 text += page.extract_text() + "\n"
# # #     except Exception as e:
# # #         raise Exception(f"Error reading PDF: {str(e)}")
# # #     return text

# # # def extract_text_from_docx(docx_path):
# # #     """Extract text from DOCX file"""
# # #     text = ""
# # #     try:
# # #         doc = Document(docx_path)
# # #         for paragraph in doc.paragraphs:
# # #             text += paragraph.text + "\n"
# # #     except Exception as e:
# # #         raise Exception(f"Error reading DOCX: {str(e)}")
# # #     return text

# # # def extract_email(text):
# # #     """Extract email address from text"""
# # #     email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
# # #     emails = re.findall(email_pattern, text)
# # #     return emails[0] if emails else ""

# # # def extract_phone(text):
# # #     """Extract phone number from text"""
# # #     phone_pattern = r'(\+\d{1,3}[-.]?)?\(?\d{3}\)?[-.]?\d{3}[-.]?\d{4}'
# # #     phones = re.findall(phone_pattern, text)
# # #     return phones[0][0] if phones else ""

# # # def extract_education(text):
# # #     """Extract education information"""
# # #     education_keywords = [
# # #         'Bachelor', 'B\.?S\.?', 'B\.?A\.?', 'Master', 'M\.?S\.?', 'M\.?A\.?', 
# # #         'PhD', 'Doctorate', 'Associate', 'Diploma', 'Certificate'
# # #     ]
    
# # #     education = []
# # #     lines = text.split('\n')
    
# # #     for line in lines:
# # #         for keyword in education_keywords:
# # #             if keyword.lower() in line.lower():
# # #                 education.append(line.strip())
# # #                 break
    
# # #     return education

# # # def extract_experience(text):
# # #     """Extract experience information and calculate years"""
# # #     # Look for years of experience patterns
# # #     exp_patterns = [
# # #         r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience',
# # #         r'experience\s*(?:of)?\s*(\d+)\+?\s*(?:years?|yrs?)',
# # #         r'(\d+)\s*(?:years?|yrs?)\s*in\s*',
# # #     ]
    
# # #     for pattern in exp_patterns:
# # #         matches = re.findall(pattern, text, re.IGNORECASE)
# # #         if matches:
# # #             try:
# # #                 return int(matches[0])
# # #             except:
# # #                 continue
    
# # #     # If no explicit years found, estimate from dates
# # #     date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4}'
# # #     dates = re.findall(date_pattern, text, re.IGNORECASE)
    
# # #     if len(dates) >= 2:
# # #         try:
# # #             years = []
# # #             for date in dates:
# # #                 year_match = re.search(r'\d{4}', date)
# # #                 if year_match:
# # #                     years.append(int(year_match.group()))
            
# # #             if years:
# # #                 max_year = max(years)
# # #                 min_year = min(years)
# # #                 return max(0, max_year - min_year)
# # #         except:
# # #             pass
    
# # #     return 0

# # # def extract_certifications(text):
# # #     """Extract certifications from text"""
# # #     cert_keywords = [
# # #         'AWS', 'Azure', 'Google Cloud', 'PMP', 'Scrum', 'ITIL', 'CISSP',
# # #         'CEH', 'CCNA', 'CCNP', 'OCP', 'MCSA', 'MCSE', 'CPA', 'CFA'
# # #     ]
    
# # #     certifications = []
# # #     for cert in cert_keywords:
# # #         if cert.lower() in text.lower():
# # #             certifications.append(cert)
    
# # #     return certifications

# # # def parse_resume(file_path):
# # #     """Main function to parse resume file"""
# # #     try:
# # #         # Determine file type and extract text
# # #         file_ext = os.path.splitext(file_path)[1].lower()
        
# # #         if file_ext == '.pdf':
# # #             text = extract_text_from_pdf(file_path)
# # #         elif file_ext in ['.docx', '.doc']:
# # #             text = extract_text_from_docx(file_path)
# # #         else:
# # #             raise ValueError(f"Unsupported file format: {file_ext}")
    
# # #         # Extract information
# # #         result = {
# # #             'text': text,
# # #             'email': extract_email(text),
# # #             'phone': extract_phone(text),
# # #             'education': extract_education(text),
# # #             'experience_years': extract_experience(text),
# # #             'certifications': extract_certifications(text)
# # #         }
        
# # #         return result
        
# # #     except Exception as e:
# # #         raise Exception(f"Failed to parse resume: {str(e)}")

# # import PyPDF2
# # from docx import Document
# # import re
# # import os

# # def extract_text_from_pdf(pdf_path):
# #     """Extract text from PDF file"""
# #     text = ""
# #     try:
# #         with open(pdf_path, 'rb') as file:
# #             pdf_reader = PyPDF2.PdfReader(file)
# #             for page in pdf_reader.pages:
# #                 text += page.extract_text() + "\n"
# #     except Exception as e:
# #         raise Exception(f"Error reading PDF: {str(e)}")
# #     return text

# # def extract_text_from_docx(docx_path):
# #     """Extract text from DOCX file"""
# #     text = ""
# #     try:
# #         doc = Document(docx_path)
# #         for paragraph in doc.paragraphs:
# #             text += paragraph.text + "\n"
# #     except Exception as e:
# #         raise Exception(f"Error reading DOCX: {str(e)}")
# #     return text

# # def extract_email(text):
# #     """Extract email address from text safely"""
# #     email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
# #     emails = re.findall(email_pattern, text)
# #     # FIX: Added check to ensure index exists before accessing
# #     if emails and len(emails) > 0:
# #         return emails[0]
# #     return ""

# # def extract_phone(text):
# #     """Extract phone number from text safely"""
# #     phone_pattern = r'(\+\d{1,3}[-.]?)?\(?\d{3}\)?[-.]?\d{3}[-.]?\d{4}'
# #     phones = re.findall(phone_pattern, text)
    
# #     # FIX: Check if matches exist and handle tuple results from capture groups
# #     if phones and len(phones) > 0:
# #         match = phones[0]
# #         if isinstance(match, tuple):
# #             # Join groups to avoid "index out of range" if a specific group is empty
# #             full_phone = "".join(match).strip()
# #             return full_phone if full_phone else ""
# #         return match
# #     return ""

# # def extract_education(text):
# #     """Extract education information"""
# #     education_keywords = [
# #         'Bachelor', 'B\.?S\.?', 'B\.?A\.?', 'Master', 'M\.?S\.?', 'M\.?A\.?', 
# #         'PhD', 'Doctorate', 'Associate', 'Diploma', 'Certificate'
# #     ]
    
# #     education = []
# #     lines = text.split('\n')
    
# #     for line in lines:
# #         for keyword in education_keywords:
# #             if keyword.lower() in line.lower():
# #                 education.append(line.strip())
# #                 break
    
# #     return education

# # def extract_experience(text):
# #     """Extract experience information and calculate years"""
# #     exp_patterns = [
# #         r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience',
# #         r'experience\s*(?:of)?\s*(\d+)\+?\s*(?:years?|yrs?)',
# #         r'(\d+)\s*(?:years?|yrs?)\s*in\s*',
# #     ]
    
# #     for pattern in exp_patterns:
# #         matches = re.findall(pattern, text, re.IGNORECASE)
# #         if matches:
# #             try:
# #                 return int(matches[0])
# #             except:
# #                 continue
    
# #     date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4}'
# #     dates = re.findall(date_pattern, text, re.IGNORECASE)
    
# #     if len(dates) >= 2:
# #         try:
# #             years = []
# #             for date in dates:
# #                 year_match = re.search(r'\d{4}', date)
# #                 if year_match:
# #                     years.append(int(year_match.group()))
            
# #             if years:
# #                 max_year = max(years)
# #                 min_year = min(years)
# #                 return max(0, max_year - min_year)
# #         except:
# #             pass
    
# #     return 0

# # def extract_certifications(text):
# #     """Extract certifications from text"""
# #     cert_keywords = [
# #         'AWS', 'Azure', 'Google Cloud', 'PMP', 'Scrum', 'ITIL', 'CISSP',
# #         'CEH', 'CCNA', 'CCNP', 'OCP', 'MCSA', 'MCSE', 'CPA', 'CFA'
# #     ]
    
# #     certifications = []
# #     for cert in cert_keywords:
# #         if cert.lower() in text.lower():
# #             certifications.append(cert)
    
# #     return certifications

# # def parse_resume(file_path):
# #     """Main function to parse resume file"""
# #     try:
# #         file_ext = os.path.splitext(file_path)[1].lower()
        
# #         if file_ext == '.pdf':
# #             text = extract_text_from_pdf(file_path)
# #         elif file_ext in ['.docx', '.doc']:
# #             text = extract_text_from_docx(file_path)
# #         else:
# #             raise ValueError(f"Unsupported file format: {file_ext}")
        
# #         result = {
# #             'text': text,
# #             'email': extract_email(text),
# #             'phone': extract_phone(text),
# #             'education': extract_education(text),
# #             'experience_years': extract_experience(text),
# #             'certifications': extract_certifications(text)
# #         }
        
# #         return result
        
# #     except Exception as e:
# #         raise Exception(f"Failed to parse resume: {str(e)}")
# # import PyPDF2
# # from docx import Document
# # import re
# # import os

# # def extract_text_from_pdf(pdf_path):
# #     """Extract text from PDF file"""
# #     text = ""
# #     try:
# #         with open(pdf_path, 'rb') as file:
# #             pdf_reader = PyPDF2.PdfReader(file)
# #             for page in pdf_reader.pages:
# #                 text += page.extract_text() + "\n"
# #     except Exception as e:
# #         raise Exception(f"Error reading PDF: {str(e)}")
# #     return text

# # def extract_text_from_docx(docx_path):
# #     """Extract text from DOCX file"""
# #     text = ""
# #     try:
# #         doc = Document(docx_path)
# #         for paragraph in doc.paragraphs:
# #             text += paragraph.text + "\n"
# #     except Exception as e:
# #         raise Exception(f"Error reading DOCX: {str(e)}")
# #     return text

# # def extract_email(text):
# #     """Extract email address from text safely"""
# #     email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
# #     emails = re.findall(email_pattern, text)
# #     if emails and len(emails) > 0:
# #         return emails[0]
# #     return ""

# # def extract_phone(text):
# #     """Extract phone number from text safely"""
# #     phone_pattern = r'(\+\d{1,3}[-.]?)?\(?\d{3}\)?[-.]?\d{3}[-.]?\d{4}'
# #     phones = re.findall(phone_pattern, text)
# #     if phones and len(phones) > 0:
# #         match = phones[0]
# #         if isinstance(match, tuple):
# #             full_phone = "".join(match).strip()
# #             return full_phone if full_phone else ""
# #         return match
# #     return ""

# # def extract_education(text):
# #     """Extract education information"""
# #     education_keywords = [
# #         'Bachelor', 'B\.?S\.?', 'B\.?A\.?', 'Master', 'M\.?S\.?', 'M\.?A\.?', 
# #         'PhD', 'Doctorate', 'Associate', 'Diploma', 'Certificate'
# #     ]
# #     education = []
# #     lines = text.split('\n')
# #     for line in lines:
# #         for keyword in education_keywords:
# #             if keyword.lower() in line.lower():
# #                 education.append(line.strip())
# #                 break
# #     return education

# # def extract_experience(text):
# #     """Extract experience information and calculate years"""
# #     exp_patterns = [
# #         r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience',
# #         r'experience\s*(?:of)?\s*(\d+)\+?\s*(?:years?|yrs?)',
# #         r'(\d+)\s*(?:years?|yrs?)\s*in\s*',
# #     ]
# #     for pattern in exp_patterns:
# #         matches = re.findall(pattern, text, re.IGNORECASE)
# #         if matches:
# #             try:
# #                 return int(matches[0])
# #             except:
# #                 continue
# #     date_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4}'
# #     dates = re.findall(date_pattern, text, re.IGNORECASE)
# #     if len(dates) >= 2:
# #         try:
# #             years = []
# #             for date in dates:
# #                 year_match = re.search(r'\d{4}', date)
# #                 if year_match:
# #                     years.append(int(year_match.group()))
# #             if years:
# #                 max_year = max(years)
# #                 min_year = min(years)
# #                 return max(0, max_year - min_year)
# #         except:
# #             pass
# #     return 0

# # def extract_certifications(text):
# #     """Extract certifications from text"""
# #     cert_keywords = [
# #         'AWS', 'Azure', 'Google Cloud', 'PMP', 'Scrum', 'ITIL', 'CISSP',
# #         'CEH', 'CCNA', 'CCNP', 'OCP', 'MCSA', 'MCSE', 'CPA', 'CFA'
# #     ]
# #     certifications = []
# #     for cert in cert_keywords:
# #         if cert.lower() in text.lower():
# #             certifications.append(cert)
# #     return certifications

# # def parse_resume(file_path):
# #     """Main function to parse resume file"""
# #     try:
# #         file_ext = os.path.splitext(file_path)[1].lower()
# #         if file_ext == '.pdf':
# #             text = extract_text_from_pdf(file_path)
# #         elif file_ext in ['.docx', '.doc']:
# #             text = extract_text_from_docx(file_path)
# #         else:
# #             raise ValueError(f"Unsupported file format: {file_ext}")
# #         return {
# #             'text': text,
# #             'email': extract_email(text),
# #             'phone': extract_phone(text),
# #             'education': extract_education(text),
# #             'experience_years': extract_experience(text),
# #             'certifications': extract_certifications(text)
# #         }
# #     except Exception as e:
# #         raise Exception(f"Failed to parse resume: {str(e)}")
# import PyPDF2
# from docx import Document
# import re
# import os

# def extract_text_from_pdf(pdf_path):
#     text = ""
#     try:
#         with open(pdf_path, 'rb') as file:
#             pdf_reader = PyPDF2.PdfReader(file)
#             for page in pdf_reader.pages:
#                 text += page.extract_text() + "\n"
#     except Exception as e:
#         raise Exception(f"Error reading PDF: {str(e)}")
#     return text

# def extract_text_from_docx(docx_path):
#     text = ""
#     try:
#         doc = Document(docx_path)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + "\n"
#     except Exception as e:
#         raise Exception(f"Error reading DOCX: {str(e)}")
#     return text

# def extract_email(text):
#     """Safely extract email"""
#     email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
#     emails = re.findall(email_pattern, text)
#     if emails and len(emails) > 0:
#         return emails[0]
#     return ""

# def extract_phone(text):
#     """Safely extract phone"""
#     phone_pattern = r'(\+\d{1,3}[-.]?)?\(?\d{3}\)?[-.]?\d{3}[-.]?\d{4}'
#     phones = re.findall(phone_pattern, text)
#     if phones and len(phones) > 0:
#         match = phones[0]
#         # Handle tuple return from regex groups
#         if isinstance(match, tuple):
#             return "".join(match).strip()
#         return match
#     return ""

# def extract_education(text):
#     keywords = ['Bachelor', 'B.S', 'B.A', 'Master', 'M.S', 'M.A', 'PhD', 'Doctorate', 'Diploma']
#     education = []
#     for line in text.split('\n'):
#         for kw in keywords:
#             if kw.lower() in line.lower():
#                 education.append(line.strip())
#                 break
#     return education

# def extract_experience(text):
#     # Regex for years of experience
#     exp_patterns = [r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience']
#     for pattern in exp_patterns:
#         matches = re.findall(pattern, text, re.IGNORECASE)
#         if matches:
#             try: return int(matches[0])
#             except: continue
            
#     # Fallback: Estimate from dates
#     date_pattern = r'\d{4}'
#     dates = re.findall(date_pattern, text)
#     if len(dates) >= 2:
#         try:
#             years = [int(d) for d in dates]
#             return max(0, max(years) - min(years))
#         except: pass
#     return 0

# def extract_certifications(text):
#     cert_keywords = ['AWS', 'Azure', 'Google Cloud', 'PMP', 'Scrum', 'CISSP', 'CCNA', 'React', 'Python']
#     found = []
#     for cert in cert_keywords:
#         if cert.lower() in text.lower():
#             found.append(cert)
#     return found

# def parse_resume(file_path):
#     try:
#         ext = os.path.splitext(file_path)[1].lower()
#         if ext == '.pdf':
#             text = extract_text_from_pdf(file_path)
#         elif ext in ['.docx', '.doc']:
#             text = extract_text_from_docx(file_path)
#         else:
#             raise ValueError("Unsupported format")
            
#         return {
#             'text': text,
#             'email': extract_email(text),
#             'phone': extract_phone(text),
#             'education': extract_education(text),
#             'experience_years': extract_experience(text),
#             'certifications': extract_certifications(text)
#         }
#     except Exception as e:
#         raise Exception(f"Failed to parse resume: {str(e)}")
import PyPDF2
from docx import Document
import re
import os

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")
    return text

def extract_email(text):
    """Safely extract email"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails and len(emails) > 0:
        return emails[0]
    return ""

def extract_phone(text):
    """Safely extract phone"""
    phone_pattern = r'(\+\d{1,3}[-.]?)?\(?\d{3}\)?[-.]?\d{3}[-.]?\d{4}'
    phones = re.findall(phone_pattern, text)
    if phones and len(phones) > 0:
        match = phones[0]
        # Handle tuple return from regex groups
        if isinstance(match, tuple):
            return "".join(match).strip()
        return match
    return ""

def extract_education(text):
    keywords = ['Bachelor', 'B.S', 'B.A', 'Master', 'M.S', 'M.A', 'PhD', 'Doctorate', 'Diploma']
    education = []
    for line in text.split('\n'):
        for kw in keywords:
            if kw.lower() in line.lower():
                education.append(line.strip())
                break
    return education

def extract_experience(text):
    # Regex for years of experience
    exp_patterns = [r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*experience']
    for pattern in exp_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            try: return int(matches[0])
            except: continue
            
    # Fallback: Estimate from dates
    date_pattern = r'\d{4}'
    dates = re.findall(date_pattern, text)
    if len(dates) >= 2:
        try:
            years = [int(d) for d in dates]
            return max(0, max(years) - min(years))
        except: pass
    return 0

def extract_certifications(text):
    cert_keywords = ['AWS', 'Azure', 'Google Cloud', 'PMP', 'Scrum', 'CISSP', 'CCNA', 'React', 'Python']
    found = []
    for cert in cert_keywords:
        if cert.lower() in text.lower():
            found.append(cert)
    return found

def parse_resume(file_path):
    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            text = extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported format")
            
        return {
            'text': text,
            'email': extract_email(text),
            'phone': extract_phone(text),
            'education': extract_education(text),
            'experience_years': extract_experience(text),
            'certifications': extract_certifications(text)
        }
    except Exception as e:
        raise Exception(f"Failed to parse resume: {str(e)}")